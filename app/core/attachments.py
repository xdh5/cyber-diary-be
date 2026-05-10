from __future__ import annotations

import csv
import json
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

MAX_EXTRACT_CHARS = 12000


def _read_text_bytes(payload: bytes) -> str:
    for encoding in ('utf-8', 'utf-8-sig', 'gb18030'):
        try:
            return payload.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return payload.decode('utf-8', errors='ignore').strip()


def _clip_text(text: str) -> str:
    normalized = text.replace('\r\n', '\n').replace('\r', '\n').strip()
    if len(normalized) <= MAX_EXTRACT_CHARS:
        return normalized
    return normalized[:MAX_EXTRACT_CHARS].rstrip() + '\n...'


def _extract_from_pdf(payload: bytes) -> str:
    reader = PdfReader(BytesIO(payload))
    parts: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ''
        if page_text.strip():
            parts.append(page_text.strip())
    return _clip_text('\n\n'.join(parts))


def _extract_from_docx(payload: bytes) -> str:
    document = Document(BytesIO(payload))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return _clip_text('\n\n'.join(parts))


def _extract_from_xlsx(payload: bytes) -> str:
    workbook = load_workbook(BytesIO(payload), data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f'[Sheet] {sheet.title}')
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if values:
                parts.append(' | '.join(values))
    return _clip_text('\n'.join(parts))


def extract_attachment_text(payload: bytes, filename: str, content_type: str | None = None) -> str:
    file_name = (filename or '').lower()
    suffix = Path(file_name).suffix.lower()
    mime_type = (content_type or '').lower()

    if suffix in {'.txt', '.md', '.csv', '.log', '.json', '.yaml', '.yml'} or mime_type.startswith('text/'):
        if suffix == '.json' or mime_type.endswith('/json'):
            try:
                data = json.loads(_read_text_bytes(payload))
                return _clip_text(json.dumps(data, ensure_ascii=False, indent=2))
            except Exception:
                return _clip_text(_read_text_bytes(payload))

        if suffix == '.csv' or mime_type.endswith('/csv'):
            try:
                decoded = _read_text_bytes(payload)
                reader = csv.reader(StringIO(decoded))
                rows = [' | '.join(cell.strip() for cell in row if cell.strip()) for row in reader]
                rows = [row for row in rows if row]
                return _clip_text('\n'.join(rows))
            except Exception:
                return _clip_text(_read_text_bytes(payload))

        return _clip_text(_read_text_bytes(payload))

    if suffix == '.pdf' or mime_type == 'application/pdf':
        return _extract_from_pdf(payload)

    if suffix == '.docx' or mime_type in {
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    }:
        return _extract_from_docx(payload)

    if suffix in {'.xlsx', '.xlsm'} or mime_type in {
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel.sheet.macroenabled.12',
    }:
        return _extract_from_xlsx(payload)

    return ''
